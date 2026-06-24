"""
Context file detection, classification, hashing, and text extraction.

Three responsibilities:

1. **Enumeration & classification**: walk --context-dir and --context-file
   inputs, classify each file as text_document | structured_data |
   unknown, hash for cache keys, assign a stable context_id (CTX-/
   DATA-/UNK-YYYYMMDD-NNN). Output: List[ContextFile].

2. **Per-format text extraction**: dispatch on file extension to the
   right reader (.md/.txt verbatim, .docx via python-docx, .pdf via
   pypdf). Returns (text, failure_reason). Structured data is detected
   but never loaded -- future skills will read it directly.

3. **Audit serialization**: ContextFile.to_audit_dict() for the
   audit.jsonl stream, ContextFile.to_state_dict() for session_state.json.
   build_context_state() rolls up all ContextFiles into the runtime_state
   payload.

This module does NO LLM calls. Summarization (which does call the
Summarizer service agent) lives in budget_manager.py. The split keeps
the read-from-disk side independent of model dispatch, which makes it
trivial to unit-test and reuse.
"""
from __future__ import annotations

import argparse
import hashlib
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, List, Optional, Set, Tuple

from adam.core.exceptions import ConfigError
from adam.context._config import _rt_context


# ============================================================
# Format catalogs
# ============================================================

# Supported text formats (loaded into the [T0] Background block)
TEXT_DOCUMENT_EXTENSIONS = {".md", ".txt", ".docx", ".pdf"}

# Detected-but-not-loaded data formats (audit-logged; future skill access)
STRUCTURED_DATA_EXTENSIONS = {".csv", ".xlsx", ".xls", ".json"}


# ============================================================
# ContextFile
# ============================================================

class ContextFile:
    """
    A single file detected by the Context Loader.

    Pass 1 captures: path, classification, hash, size, context_id.
    Pass 2 extends with: extracted_text, char_count, token_estimate,
    summary_used, summary_path, parse_status.
    """

    def __init__(
        self,
        path:         Path,
        classification: str,
        sha256:       str,
        size_bytes:   int,
        context_id:   str,
    ) -> None:
        self.path           = path
        self.classification = classification
        self.sha256         = sha256
        self.size_bytes     = size_bytes
        self.context_id     = context_id

        # Pass 2 fields - initialized to None so Pass 1 can write a
        # complete schema even though loading hasn't happened yet.
        self.parse_status:     Optional[str]  = None
        self.original_chars:   Optional[int]  = None
        self.injected_chars:   Optional[int]  = None
        self.token_estimate:   Optional[int]  = None
        self.summary_used:     Optional[bool] = None
        self.summary_path:     Optional[str]  = None
        self.failure_reason:   Optional[str]  = None

    @property
    def filename(self) -> str:
        return self.path.name

    def to_audit_dict(self) -> Dict[str, Any]:
        """Audit-log representation."""
        return {
            "context_id":     self.context_id,
            "filename":       self.filename,
            "source_path":    str(self.path),
            "classification": self.classification,
            "sha256":         self.sha256,
            "size_bytes":     self.size_bytes,
            "parse_status":   self.parse_status,
        }

    def to_state_dict(self) -> Dict[str, Any]:
        """session_state.json representation. Fields that haven't been
        populated in Pass 1 appear as null and become real values in Pass 2."""
        d: Dict[str, Any] = {
            "context_id":     self.context_id,
            "type":           self.classification,
            "filename":       self.filename,
            "source_path":    str(self.path),
            "sha256":         self.sha256,
            "size_bytes":     self.size_bytes,
            "parse_status":   self.parse_status,
        }
        if self.original_chars is not None:
            d["original_chars"] = self.original_chars
        if self.injected_chars is not None:
            d["injected_chars"] = self.injected_chars
        if self.token_estimate is not None:
            d["token_estimate"] = self.token_estimate
        if self.summary_used is not None:
            d["summary_used"] = self.summary_used
        if self.summary_path is not None:
            d["summary_path"] = self.summary_path
        if self.failure_reason is not None:
            d["failure_reason"] = self.failure_reason
        return d


# ============================================================
# Classification & hashing
# ============================================================

def _classify_file(path: Path) -> str:
    """Classify a file by extension into the three context categories."""
    ext = path.suffix.lower()
    if ext in TEXT_DOCUMENT_EXTENSIONS:
        return "text_document"
    if ext in STRUCTURED_DATA_EXTENSIONS:
        return "structured_data"
    return "unknown"


def _hash_file(path: Path) -> str:
    """
    Compute sha256 of file contents. Streamed so large files don't
    pin memory; returned as hex digest.
    """
    h = hashlib.sha256()
    with open(path, "rb") as f:
        while True:
            chunk = f.read(65536)
            if not chunk:
                break
            h.update(chunk)
    return h.hexdigest()


# ============================================================
# Enumeration & detection
# ============================================================

def _enumerate_context_files(args: argparse.Namespace) -> List[Path]:
    """
    Collect file paths from --context-dir and --context-file flags.
    Returns paths in a stable order: directory contents sorted by name,
    then per-flag --context-file entries in the order given. Duplicates
    (same resolved path) are skipped.
    """
    paths: List[Path] = []
    seen: Set[Path] = set()

    if args.context_dir is not None:
        ctx_dir = Path(args.context_dir)
        if not ctx_dir.exists():
            raise ConfigError(
                f"--context-dir path does not exist: {args.context_dir}"
            )
        if not ctx_dir.is_dir():
            raise ConfigError(
                f"--context-dir is not a directory: {args.context_dir}"
            )
        for p in sorted(ctx_dir.iterdir()):
            if p.is_file():
                resolved = p.resolve()
                if resolved not in seen:
                    paths.append(p)
                    seen.add(resolved)

    if args.context_file is not None:
        for raw in args.context_file:
            p = Path(raw)
            if not p.exists():
                raise ConfigError(f"--context-file path does not exist: {raw}")
            if not p.is_file():
                raise ConfigError(f"--context-file is not a regular file: {raw}")
            resolved = p.resolve()
            if resolved not in seen:
                paths.append(p)
                seen.add(resolved)

    return paths


def detect_context_files(args: argparse.Namespace) -> List[ContextFile]:
    """
    Enumerate, classify, hash, and assign context_id to each file
    referenced by --context-dir / --context-file.

    Returns an empty list if neither flag was used.

    Context IDs follow these formats:
      CTX-YYYYMMDD-NNN  for text_document
      DATA-YYYYMMDD-NNN for structured_data
      UNK-YYYYMMDD-NNN  for unknown (skipped but audit-logged)
    """
    paths = _enumerate_context_files(args)
    if not paths:
        return []

    today = datetime.now().strftime("%Y%m%d")
    counters: Dict[str, int] = {"text_document": 0, "structured_data": 0, "unknown": 0}
    prefix_for: Dict[str, str] = {
        "text_document":   "CTX",
        "structured_data": "DATA",
        "unknown":         "UNK",
    }

    files: List[ContextFile] = []
    for path in paths:
        cls = _classify_file(path)
        counters[cls] += 1
        cid = f"{prefix_for[cls]}-{today}-{counters[cls]:03d}"
        try:
            sha = _hash_file(path)
            size = path.stat().st_size
        except OSError as e:
            # File became unreadable between enumeration and hashing.
            # Soft-fail with a clean record rather than aborting startup.
            cf = ContextFile(path=path, classification=cls, sha256="",
                             size_bytes=0, context_id=cid)
            cf.parse_status = "io_error"
            cf.failure_reason = f"{type(e).__name__}: {e}"
            files.append(cf)
            continue
        cf = ContextFile(
            path=path, classification=cls, sha256=sha,
            size_bytes=size, context_id=cid,
        )
        files.append(cf)

    return files


def build_context_state(context_files: List[ContextFile]) -> Dict[str, Any]:
    """
    Build the context_state subtree for session_state.json from the
    detected files. Empty list -> {"enabled": False}; otherwise a
    populated structure with all files and their parse status.
    """
    if not context_files:
        return {"enabled": False, "files": []}

    counts = {"text_document": 0, "structured_data": 0, "unknown": 0,
              "io_error": 0, "extraction_failed": 0, "extracted": 0,
              "skipped_unknown_type": 0, "skipped_not_loaded": 0}
    for cf in context_files:
        counts[cf.classification] = counts.get(cf.classification, 0) + 1
        if cf.parse_status:
            counts[cf.parse_status] = counts.get(cf.parse_status, 0) + 1

    return {
        "enabled":    True,
        "file_count": len(context_files),
        "counts":     counts,
        "files":      [cf.to_state_dict() for cf in context_files],
    }


# ============================================================
# Per-format text extraction
# ============================================================

def _extract_text_md_or_txt(path: Path) -> Tuple[str, Optional[str]]:
    """Read .md or .txt verbatim. Returns (text, failure_reason)."""
    try:
        text = path.read_text(encoding="utf-8")
        return text, None
    except UnicodeDecodeError:
        try:
            text = path.read_text(encoding="latin-1")
            return text, None
        except Exception as e:
            return "", f"decode_error: {type(e).__name__}: {e}"
    except Exception as e:
        return "", f"read_error: {type(e).__name__}: {e}"


def _extract_text_docx(path: Path) -> Tuple[str, Optional[str]]:
    """
    Extract text from a .docx file using python-docx.
    Returns (text, failure_reason).
    """
    try:
        from docx import Document  # python-docx
    except ImportError:
        return "", "python-docx not installed (should have been caught at startup)"

    try:
        doc = Document(str(path))
    except Exception as e:
        return "", f"docx_open_error: {type(e).__name__}: {e}"

    parts: List[str] = []
    # Paragraphs
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            parts.append(text)
    # Table cells -- read in row-major order
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text = cell.text.strip()
                if text:
                    parts.append(text)

    full = "\n".join(parts)
    if not full.strip():
        return "", "docx_empty_or_unreadable"
    return full, None


def _extract_text_pdf(path: Path, min_chars: int, scan_threshold_bytes: int) -> Tuple[str, Optional[str]]:
    """
    Extract text from a PDF using pypdf. Heuristics:
      - If pypdf isn't installed, fail soft.
      - If the file is large but extracted text is small, the PDF is
        likely scanned (image-only) and not usable without OCR.
    Returns (text, failure_reason).
    """
    try:
        from pypdf import PdfReader
    except ImportError:
        return "", "pypdf not installed (should have been caught at startup)"

    try:
        reader = PdfReader(str(path))
    except Exception as e:
        return "", f"pdf_open_error: {type(e).__name__}: {e}"

    parts: List[str] = []
    for page in reader.pages:
        try:
            page_text = page.extract_text() or ""
        except Exception:
            page_text = ""
        page_text = page_text.strip()
        if page_text:
            parts.append(page_text)

    full = "\n".join(parts)
    size = path.stat().st_size
    if len(full) < min_chars and size > scan_threshold_bytes:
        return "", f"pdf_likely_scanned (size={size}b, extracted_chars={len(full)})"
    if not full.strip():
        return "", "pdf_empty_or_unreadable"
    return full, None


def extract_text_for_file(cf: ContextFile) -> Tuple[str, Optional[str]]:
    """
    Dispatch to the right extractor for a text_document file.
    Returns (text, failure_reason). For non-text-document files,
    returns ("", "not a text document").
    """
    if cf.classification != "text_document":
        return "", "not a text document"

    ext = cf.path.suffix.lower()
    if ext in (".md", ".txt"):
        return _extract_text_md_or_txt(cf.path)
    if ext == ".docx":
        return _extract_text_docx(cf.path)
    if ext == ".pdf":
        return _extract_text_pdf(
            cf.path,
            min_chars=_rt_context("pdf_min_text_chars"),
            scan_threshold_bytes=_rt_context("pdf_likely_scanned_size_bytes"),
        )
    return "", f"unsupported text-document extension: {ext}"
